"""Word (.docx) export of one conversation's learnt notes (vocabulary/
recurring mistakes + lesson log) - not a full transcript. Backs the
Settings modal's Data controls tab (see routes_api.export_notes_docx and
settings.js). Builds the document entirely in memory; the FastAPI route
hands the raw bytes straight back as a download.
"""

from __future__ import annotations

import io
from datetime import UTC, datetime

from docx import Document
from docx.shared import Pt


def _format_ts(value) -> str:
    if not value:
        return ""
    # vocab_mistakes timestamps are unix seconds; lesson_log timestamps are
    # ISO strings (see memory.py) - handle both, same split notes.js
    # already does client-side.
    try:
        if isinstance(value, (int, float)):
            dt = datetime.fromtimestamp(value, tz=UTC)
        else:
            dt = datetime.fromisoformat(str(value))
        return dt.strftime("%b %d, %Y")
    except (ValueError, OSError, TypeError):
        return str(value)


def build_notes_docx(
    profile_name: str,
    language_label: str,
    vocab_mistakes: list[dict],
    lesson_log: list[dict],
) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading(f"{language_label} - Learnt Notes", level=1)
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"{profile_name} - exported {datetime.now(tz=UTC).strftime('%b %d, %Y')}")
    subtitle_run.italic = True

    doc.add_heading("Vocabulary & Mistakes", level=2)
    if not vocab_mistakes:
        doc.add_paragraph("Nothing tracked yet.")
    else:
        for item in vocab_mistakes:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(item.get("term", ""))
            run.bold = True
            note = item.get("note")
            if note:
                p.add_run(f" - {note}")
            meta = f" (seen {item.get('occurrences', 1)}x, last {_format_ts(item.get('last_seen_ts'))})"
            meta_run = p.add_run(meta)
            meta_run.italic = True
            meta_run.font.size = Pt(9)

    doc.add_heading("Lesson Log", level=2)
    if not lesson_log:
        doc.add_paragraph("No lesson log entries yet.")
    else:
        for entry in lesson_log:
            p = doc.add_paragraph()
            date_run = p.add_run(f"{_format_ts(entry.get('ts'))}: ")
            date_run.bold = True
            p.add_run(entry.get("summary", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
